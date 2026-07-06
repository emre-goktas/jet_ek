import io
import json
import base64
import logging
from datetime import datetime
from pathlib import Path
# pyrefly: ignore [missing-import]
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

TEMPLATES_FILE = Path(__file__).parent.parent / "data" / "templates.json"

_templates_cache: list[dict] | None = None
_templates_mtime: float | None = None

def get_templates() -> list[dict]:
    """Returns the parsed templates.json, cached in memory and refreshed
    automatically whenever the file's mtime changes (so manual edits are picked
    up without a server restart, but repeated calls skip the disk read)."""
    global _templates_cache, _templates_mtime

    if not TEMPLATES_FILE.exists():
        return []

    mtime = TEMPLATES_FILE.stat().st_mtime
    if _templates_cache is not None and mtime == _templates_mtime:
        return _templates_cache

    with open(TEMPLATES_FILE, "r", encoding="utf-8") as f:
        _templates_cache = json.load(f)
    _templates_mtime = mtime
    return _templates_cache

def number_to_turkish_words(n):
    if n == 0:
        return "sıfır"

    ones = ["", "bir", "iki", "üç", "dört", "beş", "altı", "yedi", "sekiz", "dokuz"]
    tens = ["", "on", "yirmi", "otuz", "kırk", "elli", "altmış", "yetmiş", "seksen", "doksan"]
    scales = ["", "bin", "milyon", "milyar", "trilyon"]

    if n >= 10 ** (3 * len(scales)):
        raise ValueError(f"Number too large to convert to Turkish words: {n}")

    def read_three_digits(num):
        words = []
        h = num // 100
        t = (num % 100) // 10
        o = num % 10

        if h > 1:
            words.append(ones[h])
        if h > 0:
            words.append("yüz")

        if t > 0:
            words.append(tens[t])

        if o > 0:
            words.append(ones[o])

        return "".join(words)

    result = []
    scale_idx = 0

    while n > 0:
        chunk = n % 1000
        if chunk > 0:
            chunk_words = read_three_digits(chunk)

            # 1 bin denmez, sadece bin denir.
            if scale_idx == 1 and chunk == 1:
                result.append("bin")
            else:
                result.append(chunk_words + scales[scale_idx])

        n //= 1000
        scale_idx += 1

    return "".join(reversed(result))

def _set_cell_text(cell, text, align="left", bold=False):
    p = cell.paragraphs[0]
    p.text = text
    p.runs[0].font.name = 'Times New Roman'
    p.runs[0].font.size = Pt(12)
    if bold:
        p.runs[0].font.bold = True

    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

def generate_docx_from_template(files_data: list[dict]) -> bytes:
    templates = get_templates()
    if not templates:
        raise ValueError("No templates available.")

    template = templates[0]

    if template.get("file_path"):
        return _generate_from_docx_template(template, files_data)

    return _generate_legacy_docx(template, files_data)


def _generate_from_docx_template(template: dict, files_data: list[dict]) -> bytes:
    """Handles 'file_path' mode: fills in an existing .docx template using native python-docx."""
    doc = docx.Document(template["file_path"])

    start_num = files_data[0].get("ek_no", 1) if files_data else 1
    end_num = files_data[-1].get("ek_no", 1) if files_data else 1
    total_pages = sum(f.get("page_count", 0) for f in files_data)

    if doc.tables:
        t = doc.tables[0]
        if not t.rows:
            raise ValueError("Template table has no rows; expected at least a TOPLAM row.")

        toplam_row = t.rows[-1]
        toplam_tr = toplam_row._tr

        # Put total pages in the correct cell (Sayfa Adedi column is index 2)
        if len(toplam_row.cells) > 2:
            _set_cell_text(toplam_row.cells[2], str(total_pages), align="center")

        # Remove all empty placeholder rows (from index 1 to second-to-last)
        for row in t.rows[1:-1]:
            t._tbl.remove(row._tr)

        # Insert data rows
        for f in files_data:
            new_row = t.add_row()
            _set_cell_text(new_row.cells[0], str(f.get("ek_no", "")), align="center")
            _set_cell_text(new_row.cells[1], str(f.get("mahiyet", "")), align="left")
            _set_cell_text(new_row.cells[2], str(f.get("page_count", "")), align="center")
            _set_cell_text(new_row.cells[3], "", align="center")
            _set_cell_text(new_row.cells[4], "F", align="center")

        # Move TOPLAM row to the end
        t._tbl.append(toplam_tr)

    # Replace paragraph placeholders ( )
    for p in doc.paragraphs:
        if "( ) numaraları altında" in p.text:
            p.text = p.text.replace("( ) numaraları", f"({start_num}-{end_num}) numaraları")
            p.text = p.text.replace("( ) sayfadan", f"({total_pages}) sayfadan")

    out_stream = io.BytesIO()
    doc.save(out_stream)
    return out_stream.getvalue()


def _generate_legacy_docx(template: dict, files_data: list[dict]) -> bytes:
    """Handles the manual python-docx mode (e.g. saglik_bakanligi): builds the
    whole document from scratch based on the template's JSON configuration."""
    doc = docx.Document()
    total_pages = sum(f.get("page_count", 0) for f in files_data)

    _add_logo(doc, template)
    _add_header(doc, template)
    doc.add_paragraph()  # Spacer

    table = _build_table(doc, template, files_data)
    _add_footer_rows(doc, table, template, total_pages)
    _add_post_table_text(doc, template, files_data, total_pages)
    _add_signature_block(doc, template)

    out_stream = io.BytesIO()
    doc.save(out_stream)
    return out_stream.getvalue()


def _add_logo(doc, template: dict) -> None:
    if not template.get("logo_base64"):
        return
    try:
        image_bytes = base64.b64decode(template["logo_base64"])
        image_stream = io.BytesIO(image_bytes)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(image_stream, width=Inches(1.0))
    except Exception as e:
        logging.warning(f"Logo embed failed, continuing without logo: {e}")


def _add_header(doc, template: dict) -> None:
    if not template.get("header_text"):
        return
    p_header = doc.add_paragraph()
    p_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line in template["header_text"].split("\n"):
        if line == "---":
            run = p_header.add_run("___________________________________________________________\n")
            run.bold = True
        else:
            run = p_header.add_run(line + "\n")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.bold = True


def _build_table(doc, template: dict, files_data: list[dict]):
    table_config = template.get("table", {})
    header_rows = table_config.get("header_rows", [])
    data_columns = table_config.get("data_columns", [])

    # Calculate total columns from data_columns
    num_cols = len(data_columns)
    table = doc.add_table(rows=0, cols=num_cols)
    table.style = 'Table Grid'

    # Build Headers — add all rows first, then merge
    header_cells = []
    for _ in range(len(header_rows)):
        header_cells.append(table.add_row().cells)

    # Apply text and merging
    for r_idx, row_def in enumerate(header_rows):
        c_idx = 0
        for cell_def in row_def:
            # Skip cells that were already covered by a rowspan from above
            while c_idx < num_cols and header_cells[r_idx][c_idx].text != "":
                c_idx += 1

            if c_idx >= num_cols:
                break

            colspan = cell_def.get("colspan", 1)
            rowspan = cell_def.get("rowspan", 1)

            top_left_cell = header_cells[r_idx][c_idx]
            _set_cell_text(top_left_cell, cell_def.get("text", ""), align="center", bold=True)

            # Merge horizontally
            if colspan > 1:
                bottom_right_cell = header_cells[r_idx][c_idx + colspan - 1]
                top_left_cell.merge(bottom_right_cell)

            # Merge vertically
            if rowspan > 1:
                for r_offset in range(1, rowspan):
                    bottom_cell = header_cells[r_idx + r_offset][c_idx]
                    top_left_cell.merge(bottom_cell)

            c_idx += colspan

    # Add Data Rows
    for f in files_data:
        row_cells = table.add_row().cells
        for i, col_def in enumerate(data_columns):
            ctype = col_def.get("type")
            text = ""
            if ctype == "ek_no":
                text = str(f.get("ek_no", ""))
            elif ctype == "sayfa_sayisi":
                text = str(f.get("page_count", ""))
            elif ctype == "mahiyet":
                text = f.get("mahiyet", "")
            elif ctype == "constant":
                text = col_def.get("value", "")

            align = "center" if ctype in ["ek_no", "sayfa_sayisi", "constant"] else "left"
            _set_cell_text(row_cells[i], text, align=align)

    return table


def _add_footer_rows(doc, table, template: dict, total_pages: int) -> None:
    footer_rows = template.get("footer_rows", [])
    if not footer_rows:
        return

    total_pages_words = number_to_turkish_words(total_pages).capitalize()
    for row_def in footer_rows:
        row_cells = table.add_row().cells
        c_idx = 0
        for cell_def in row_def:
            colspan = cell_def.get("colspan", 1)
            text = cell_def.get("text", "")
            text = text.replace("{toplam_sayfa}", str(total_pages))
            text = text.replace("{toplam_sayfa_yazi_ile}", total_pages_words)

            top_left_cell = row_cells[c_idx]
            _set_cell_text(top_left_cell, text, align=cell_def.get("align", "left"), bold=cell_def.get("bold", False))

            if colspan > 1:
                bottom_right_cell = row_cells[c_idx + colspan - 1]
                top_left_cell.merge(bottom_right_cell)

            c_idx += colspan


def _add_post_table_text(doc, template: dict, files_data: list[dict], total_pages: int) -> None:
    if not template.get("post_table_text"):
        return

    doc.add_paragraph()
    start_num = files_data[0].get("ek_no", 1) if files_data else 1
    end_num = files_data[-1].get("ek_no", 1) if files_data else 1

    text = template["post_table_text"]
    text = text.replace("{gunun_tarihi}", datetime.now().strftime("%d.%m.%Y"))
    text = text.replace("{baslangic_no}", str(start_num))
    text = text.replace("{bitis_no}", str(end_num))
    text = text.replace("{toplam_sayfa}", str(total_pages))

    p = doc.add_paragraph()
    if template.get("post_table_align") == "justify":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = template.get("post_table_bold", False)


def _add_signature_block(doc, template: dict) -> None:
    if not template.get("post_table_signature"):
        return

    doc.add_paragraph()  # Spacer
    doc.add_paragraph()  # Spacer
    sig = template["post_table_signature"]

    # No borders, so it just reads as aligned text.
    sig_table = doc.add_table(rows=2, cols=2)

    c1 = sig_table.cell(0, 1)
    _set_cell_text(c1, sig.get("name", ""), align="center", bold=False)

    c2 = sig_table.cell(1, 1)
    _set_cell_text(c2, sig.get("title", ""), align="center", bold=False)
